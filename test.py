#!/usr/bin/env ./autotype/bin/python

# -*- coding: utf-8 -*-
# @Author: KICC
# @Date:   2018-09-17 14:31:54
# @Last Modified by:   KICC
# @Last Modified time: 2018-09-17 14:36:16

from docx import Document
from PIL import Image, ImageDraw
from io import BytesIO
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn


d1 = Document('demo1.docx')
d2 = Document('demo2.docx')

f1 = d1.paragraphs[1].paragraph_format.first_line_indent
f3 = d1.paragraphs[1].paragraph_format.left_indent
f5 = d1.paragraphs[1].paragraph_format.right_indent
f7 = d1.paragraphs[1].paragraph_format.space_before

f2 = d2.paragraphs[1].paragraph_format.first_line_indent
f4 = d2.paragraphs[1].paragraph_format.left_indent
f6 = d2.paragraphs[1].paragraph_format.right_indent
f8 = d2.paragraphs[1].paragraph_format.space_before

print(f1.pt)
print(f3)
print(f5)
print(f7)

print(f2.pt)
print(f4)
print(f6)
print(f8)
