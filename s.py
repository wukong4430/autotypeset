#!/usr/bin/env ./autotype/bin/python

# -*- coding: utf-8 -*-
# @Author: KICC
# @Date:   2018-07-27 12:03:41
# @Last Modified by:   KICC
# @Last Modified time: 2018-09-18 13:36:13

from docx import Document
from PIL import Image, ImageDraw
from io import BytesIO
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn

from utils import *


def draw_circles():
    document = Document()
    p = document.add_paragraph()
    r = p.add_run()
    img_size = 20
    for x in range(255):
        im = Image.new("RGB", (img_size, img_size), "white")
        draw_obj = ImageDraw.Draw(im)
        draw_obj.ellipse((0, 0, img_size - 1, img_size - 1),
                         fill=255 - x)  # 画圈
        fake_buf_file = BytesIO()  # 用BytesIO将图片保存在内存里, 减少磁盘操作
        im.save(fake_buf_file, 'png')
        r.add_picture(fake_buf_file)
        fake_buf_file.close()
        document.save('circle.docx')


def paragraph_operation():
    """
    Block-level item

    """
    document = Document('demo1.docx')
    p = document.paragraphs

    print(p[0].text)
    print(p[1].text)
    print(p[2].text)
    print(p[3].text)
    # get style from p[0]

    print(p[1].text[:10])
    print(p[1].style)

    # 设置行间距的格式(固定值, 多倍行距..)
    p[1].paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    # 修改行间距, 默认非固定值行距, 多倍行距, 需要根据字体大小设置
    # 用Pt() 后就是固定值的行距.
    p[1].paragraph_format.line_spacing = Pt(20)

    # 在段落结尾添加text
    # p[0].add_run(text='I wanna have a try')
    print("Line spacing is:", p[1].paragraph_format.line_spacing.pt)

    # 段落第一行的缩进, 设置完, 没有改变(原因未知)
    p[1].paragraph_format.first_line_indent = Pt(-50)
    print("First line indent is:", p[1].paragraph_format.first_line_indent.pt)

    # if True: 单一段落不会落入到两个page中
    p[1].paragraph_format.keep_together = True
    print(p[1].paragraph_format.keep_together)

    # 整个段落最左边distance 页面左侧的distance
    print(p[1].paragraph_format.left_indent)
    p[1].paragraph_format.left_indent = 266700
    print(p[1].paragraph_format.left_indent.pt)

    # 本段落与上一段落之间的distance
    print(p[1].paragraph_format.space_before)
    p[1].paragraph_format.left_indent = 266700
    print(p[1].paragraph_format.space_before)

    # 本段落与下一段落的distance
    print(p[1].paragraph_format.space_before)
    p[1].paragraph_format.left_indent = 266700
    print(p[1].paragraph_format.space_before)

    # document.save('demo1.docx')


def character_operation():
    """
    Character-level item.

    """
    document = Document('demo1.docx')
    # print('paragraph 2:', document.paragraphs[1].text)
    all_paragraphs = document.paragraphs

    for idx, paragraph in enumerate(all_paragraphs):
        if paragraph.text == '':
            continue
        runs = paragraph.runs
        if len(runs) == 1:
            # 整个段落只有一个run
            back_text = runs[0].text
            back_format = paragraph.paragraph_format
            print('是哪个段落:', paragraph.text)
            print('indent :', back_format.first_line_indent)
            if idx == len(all_paragraphs) - 1:
                # 如果是最后一个段落
                # p = all_paragraphs[idx-1]
                p = document.add_paragraph()

            else:
                p = all_paragraphs[idx + 1].insert_paragraph_before()

            runs[0].clear()
            # 改变font
            modify_character(back_text=back_text, p=p,
                             category='正文')
            # 调整段落style
            adjust_paragraph_style(formats=back_format,
                                   paragraph=p, category='正文')
            # 删除原本的段落
            delete_paragraph(paragraph)

        else:
            # 整个段落不只有一个run
            print("该段落不止一个run, 该段落是:")
            print(paragraph.text)

    # paragraph_modify = document.paragraphs[1]
    # runs = paragraph_modify.runs
    # back_text = runs[0].text
    # print(back_text)

    # p = document.paragraphs[2].insert_paragraph_before()
    # runs[0].clear()
    # for i, ch in enumerate(back_text):
    #     run = p.add_run(ch)
    #     font = run.font
    #     font.size = Pt(10)
    #     if ch in ['1', '2', '3']:
    #         font.name = u'微软雅黑'
    #         run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    #     else:
    #         font.name = u'宋体'
    #         run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # # 将原本的paragraph_format应用到生成的新的paragraph上
    # # TODO
    # # paragraph.style = origin_style
    # print("After Modifying :", p.style.font.name)

    # adjust_paragraph_style(paragraph=p)

    # delete_paragraph(paragraph_modify)
    # document.save('demo1.docx')


def section_operation():
    document = Document('demo1.docx')

    s = document.sections[-1]

    # 直接修改页边距
    # s[0].top_margin = 1514400
    # s[0].orientation = WD_ORIENT.LANDSCAPE  # index 1; another index 0
    # 页边距
    print(s.bottom_margin)
    print(s.top_margin)
    print(s.left_margin)
    print(s.right_margin)

    print(s.footer_distance)  # 页脚distance
    print(s.header_distance)  # yemei distance
    print(s.orientation)

    document.save('demo1.docx')


def main():
    character_operation()
    # adjust_paragraph_style()


if __name__ == '__main__':
    main()
